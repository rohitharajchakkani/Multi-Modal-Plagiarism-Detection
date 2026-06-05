"""
PlagiarismAI - Text Extractor module.
Supports TXT, PDF, DOC, and DOCX document extraction.
"""
import os
import re

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import docx
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


def extract_txt(file_path):
    """Read a standard plain text file."""
    for encoding in ('utf-8', 'latin-1', 'cp1252', 'utf-16'):
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    # If all fail, read as binary and ignore errors
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def extract_docx(file_path):
    """Extract text from DOCX files using python-docx."""
    if not PYTHON_DOCX_AVAILABLE:
        return "Error: python-docx is not installed. Run 'pip install python-docx' to support Word files."
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error reading DOCX file: {str(e)}"


def extract_doc_legacy(file_path):
    """
    Extract text from legacy binary DOC format.
    Uses pure-python printable strings regex extraction to avoid heavy COM/antiword dependencies.
    """
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
        
        # Extract ASCII sequences that resemble printable sentences
        # Matches runs of printable text of length 6 or more
        ascii_matches = re.findall(rb'[a-zA-Z0-9\s\.,;:!?\(\)\-\"\'\r\n\t]{6,}', content)
        decoded_blocks = []
        for block in ascii_matches:
            # Clean control chars
            decoded = block.decode('ascii', errors='ignore').strip()
            # If the block has some words, keep it
            if len(decoded.split()) > 1:
                decoded_blocks.append(decoded)
                
        cleaned_text = "\n".join(decoded_blocks)
        
        # Filter out extreme binary remnants
        cleaned_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', cleaned_text)
        return cleaned_text
    except Exception as e:
        return f"Error parsing legacy DOC: {str(e)}"


def extract_pdf(file_path):
    """
    Extract text from PDF using pdfplumber, falling back to PyMuPDF.
    If it is a scanned document (very little text extracted), we run OCR if pytesseract is available.
    """
    text = ""
    # Try pdfplumber first
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(file_path) as pdf:
                pages_text = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
                text = "\n".join(pages_text)
        except Exception:
            text = ""

    # Fallback to PyMuPDF
    if not text.strip() and PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(file_path)
            pages_text = []
            for page in doc:
                page_text = page.get_text()
                if page_text:
                    pages_text.append(page_text)
            text = "\n".join(pages_text)
        except Exception as e:
            if not text:
                return f"Error parsing PDF: {str(e)}"

    # If it is a scanned PDF (text content is extremely short)
    # Let's import the OCR extractor from ocr_engine to run OCR on the pages
    if len(text.strip()) < 50:
        from engine.ocr_engine import extract_text_from_pdf
        ocr_text = extract_text_from_pdf(file_path)
        if ocr_text and not ocr_text.startswith("Error"):
            text = ocr_text

    return text


def extract_text_from_document(file_path):
    """
    Route the file extraction to the appropriate parser based on file extension.
    """
    if not os.path.exists(file_path):
        return f"Error: File does not exist at {file_path}"

    ext = file_path.lower().split('.')[-1]
    
    if ext == 'txt':
        return extract_txt(file_path)
    elif ext == 'docx':
        return extract_docx(file_path)
    elif ext == 'doc':
        return extract_doc_legacy(file_path)
    elif ext == 'pdf':
        return extract_pdf(file_path)
    elif ext in ['jpg', 'jpeg', 'png', 'bmp', 'tiff']:
        from engine.ocr_engine import extract_text_from_file
        return extract_text_from_file(file_path)
    else:
        # Fallback: treat as plain text if it is readable
        try:
            return extract_txt(file_path)
        except Exception:
            return f"Unsupported file type: .{ext}"
